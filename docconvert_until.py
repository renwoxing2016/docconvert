# -*- coding: utf-8 -*-
"""
Created on Mon Feb 26 16:00:44 2018

@author: 
"""

# #######################
# #此处理包括如下几种文件转换
# #docx --> txt
# #docx --> xml
# #doc --> txt
# #
# #http://blog.csdn.net/g0ose/article/details/64538787
# #

from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn


# #######################
# #此处理 docx --> txt

# #
# #新建docx文档
# #
def create_docx(path):
    # #打开文档
    document = Document()
    # #加入不同等级的标题
    document.add_heading(u'MS WORD写入测试',0)
    document.add_heading(u'一级标题',1)
    document.add_heading(u'二级标题',2)
    # #添加文本
    paragraph = document.add_paragraph(u'我们在做文本测试！')
    # #设置字号
    run = paragraph.add_run(u'设置字号、')
    run.font.size = Pt(24)
    
    # #设置字体
    run = paragraph.add_run('Set Font,')
    run.font.name = 'Consolas'
    
    # #设置中文字体
    run = paragraph.add_run(u'设置中文字体、')
    run.font.name=u'宋体'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    
    # #设置斜体
    run = paragraph.add_run(u'斜体、')
    run.italic = True
    
    # #设置粗体
    run = paragraph.add_run(u'粗体').bold = True
    
    # #增加引用
    document.add_paragraph('Intense quote', style='Intense Quote')
    
    # #增加无序列表
    document.add_paragraph(u'无序列表元素1', style='List Bullet')
    document.add_paragraph(u'无序列表元素2', style='List Bullet')
    # #增加有序列表
    document.add_paragraph(u'有序列表元素1', style='List Number')
    document.add_paragraph(u'有序列表元素2', style='List Number')
    # #增加图像（此处用到图像image.bmp，请自行添加脚本所在目录中）
    #document.add_picture('image.bmp', width=Inches(1.25))
    document.add_picture('image.jpg', width=Inches(1.25))
    
    # #增加表格
    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    # #再增加3行表格元素
    for i in range(3):
        row_cells = table.add_row().cells
        row_cells[0].text = 'test'+str(i)
        row_cells[1].text = str(i)
        row_cells[2].text = 'desc'+str(i)
    
    # #增加分页
    document.add_page_break()
    
    # #保存文件
    document.save(path)

# #
# #读取docx文档 转换为txt
# #参数 docx的文件名 含路径
# #返回值 string数组
def convert_docx_2_text(path):
    # 打开文档
    docfile = Document(path)
    # 段落数 每个回车隔离一段
    #section_number = str(len(docfile.paragraphs))
    #print("段落数:"+section_number)
    
    # 输出每一段的内容
    #for para in docfile.paragraphs:
    #    print(para.text)
    
    # 输出段落编号及段落内容
    #for i in range(len(docfile.paragraphs)):
    #    # 段落内容
    #    strcontext = docfile.paragraphs[i].text
    #    print("第"+str(i)+"段的内容是："+strcontext)
    
    # 读取每段资料
    l = [ paragraph.text for paragraph in docfile.paragraphs];
    # 输出并观察结果，也可以通过其他手段处理文本即可
    #for i in l:
    #    print(i)
    
    #l = [ paragraph.text.encode('gb2312') for paragraph in docfile.paragraphs];
    
    # 读取表格材料，并输出结果
    #tables = [table for table in docfile.tables];
    #for table in tables:
    #    for row in table.rows:
    #        for cell in row.cells:
    #            print(cell.text.encode('gb2312'),'\t')
    #        print('')
    #    print('\n')
    return l


# #######################
# #此处理 docx --> xml
# #目前仅支持对url的pdf文档
# #把一个Word 文档读成一个二进制文件对象 BytesIO与StringIO类似
# #再用 Python 的标准库 zipfile 解压（所有的 .docx 文件为了节省空间都进行过压缩）
# #然后读取这个解压文件，就变成 XML了

from zipfile import ZipFile
from urllib.request import urlopen
from io import BytesIO
from bs4 import BeautifulSoup

# #参数 docx的URL文件名 含路径 
# #返回值 string数组
def convert_docx_2_xml(path):
    # #得确认是否必须是url的文件名
    wordFile = urlopen(path).read()
    #wordFile = open(path, 'a')
    wordFile = BytesIO(wordFile)
    document = ZipFile(wordFile)
    xml_content = document.read('word/document.xml')
    wordObj = BeautifulSoup(xml_content.decode('utf-8'))
    textStrings = wordObj.findAll("w:t")
    for textElem in textStrings:
        print(textElem.text) 
    return textStrings

# #######################
# #此处理 doc --> txt


# #######################
# #ceshi

#convert_docx_2_xml(u'201802-08信息.docx') #ng

#doctext = convert_docx_2_text(u'201802-08信息.docx')
#create_docx(u'测试.docx')


